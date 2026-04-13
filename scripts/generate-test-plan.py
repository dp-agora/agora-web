import sys
sys.path.insert(0, '.pip')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=18, bold=True, color=(31, 56, 100))
    elif level == 2:
        set_font(run, size=14, bold=True, color=(31, 56, 100))
    elif level == 3:
        set_font(run, size=12, bold=True, color=(70, 90, 120))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        set_font(r1, bold=True)
    r2 = p.add_run(text)
    set_font(r2)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('Note: ' + text)
    set_font(run, size=10, color=(120, 120, 120))
    p.paragraph_format.space_after = Pt(6)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], '1F3864')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                set_font(run, size=10, bold=True, color=(255, 255, 255))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            bg = 'F0F4FA' if r_idx % 2 == 0 else 'FFFFFF'
            shade_cell(row_cells[c_idx], bg)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    set_font(run, size=10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])

    doc.add_paragraph()
    return table

# ── Cover Page ───────────────────────────────────────────────────────────────

doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('SEO & AEO Implementation\nTest Plan')
set_font(title_run, size=26, bold=True, color=(31, 56, 100))
title_p.paragraph_format.space_before = Pt(60)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Ágora — agoralatam.com\nApril 2026')
set_font(sub_run, size=13, color=(100, 100, 100))

doc.add_paragraph()
sep_p = doc.add_paragraph()
sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sep_run = sep_p.add_run('─' * 55)
set_font(sep_run, size=11, color=(180, 180, 180))

doc.add_page_break()

# ── Overview ─────────────────────────────────────────────────────────────────

add_heading(doc, 'Overview', 1)
add_body(doc, (
    'This document defines the testing protocol for all SEO and AEO changes implemented on the '
    'cleanup/seo-fixes branch. Tests are grouped into two categories: '
    'Local Browser Tests (run against the dev server before merging) and '
    'Offline Discoverability Checks (run post-deploy using Google, LLMs, and free tools). '
    'Run Part 1 before opening the PR. Run Part 2 after deploying to production.'
))

add_heading(doc, 'Changes in Scope', 2)
changes = [
    'Canonical tags fixed on ~12 pages (previously all inherited canonical: /)',
    'hreflang / locale-aware metadata set on all key pages',
    'sitemap.ts updated to include all insight articles (EN + ES) with alternates',
    'Article JSON-LD schema added to insight article pages',
    'Person JSON-LD schema added to team member profile pages',
    'FAQPage JSON-LD + visible FAQ accordion sections added to homepage and /about',
    'Organization schema enriched with awards and directory sameAs links',
    'HomeFAQ and AboutPage.faq i18n keys added (EN + ES translations)',
    '301 redirects: /services → /practices and /es/services → /es/practices',
]
for c in changes:
    add_bullet(doc, c)

# ── Part 1 ───────────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, 'Part 1 — Local Browser Tests', 1)
add_note(doc, 'Prerequisites: run npm run dev. All URLs use http://localhost:3000. DevTools = F12 / Cmd+Option+I.')

# 1.1 Canonicals
add_heading(doc, '1.1  Canonical Tags', 2)
add_body(doc, (
    'For each URL below, open DevTools → Elements → search the <head> for '
    '"rel=\\"canonical\\"". The href must match the expected value exactly. '
    'No page should show canonical: /.'
))

add_table(doc,
    headers=['URL (localhost:3000)', 'Expected canonical href'],
    rows=[
        ['/', 'https://www.agoralatam.com/'],
        ['/es', 'https://www.agoralatam.com/es'],
        ['/about', 'https://www.agoralatam.com/about'],
        ['/es/about', 'https://www.agoralatam.com/es/about'],
        ['/contact', 'https://www.agoralatam.com/contact'],
        ['/careers', 'https://www.agoralatam.com/careers'],
        ['/team', 'https://www.agoralatam.com/team'],
        ['/practices', 'https://www.agoralatam.com/practices'],
        ['/practices/investment-arbitration', 'https://www.agoralatam.com/practices/investment-arbitration'],
        ['/privacy-policy', 'https://www.agoralatam.com/privacy-policy'],
        ['/legal-terms', 'https://www.agoralatam.com/legal-terms'],
        ['/disclaimers', 'https://www.agoralatam.com/disclaimers'],
    ],
    col_widths=[7.5, 9.5]
)

# 1.2 hreflang
add_heading(doc, '1.2  hreflang Tags', 2)
add_body(doc, 'On the homepage (/) and /about, confirm three <link rel="alternate"> tags exist in <head>:')
add_bullet(doc, 'hreflang="en"  →  https://www.agoralatam.com/')
add_bullet(doc, 'hreflang="es"  →  https://www.agoralatam.com/es')
add_bullet(doc, 'hreflang="x-default"  →  https://www.agoralatam.com/')
add_body(doc, 'Also check that /es and /es/about have these same three tags present.')

# 1.3 FAQ Sections
add_heading(doc, '1.3  FAQ Sections — Visual & Functional', 2)
add_body(doc, 'Test both / (homepage) and /about in EN and ES (4 URLs total):')
add_bullet(doc, 'Scroll down — the FAQ section must be visible without any JS interaction')
add_bullet(doc, 'Click each accordion item — it must expand/collapse with a chevron rotation')
add_bullet(doc, 'All 6 questions must render; verify EN and ES copy differs correctly')
add_bullet(doc, 'Confirm the section title renders (not a missing translation key like "HomeFAQ.title")')

# 1.4 Redirects
add_heading(doc, '1.4  301 Redirects — /services', 2)
add_body(doc, 'Navigate to each URL and verify the redirect chain in DevTools → Network tab:')
add_bullet(doc, '/services  →  should land at /practices with a 301 status on the first request')
add_bullet(doc, '/es/services  →  should land at /es/practices with a 301 status')
add_body(doc, 'Filter the Network tab by "services" to isolate the response. The final page must be /practices, not a blank or error page.')

# 1.5 JSON-LD
add_heading(doc, '1.5  JSON-LD Schema Validation', 2)
add_body(doc, (
    'For each page type, open DevTools → Elements → search for '
    '"application/ld+json". Copy the JSON content and paste into '
    'https://validator.schema.org or the Google Rich Results Test '
    '(you can paste raw HTML source for localhost pages).'
))

add_table(doc,
    headers=['Page', 'Expected Schema Type(s)', 'Key Fields to Spot-Check'],
    rows=[
        ['/ (homepage)', 'FAQPage + Organization + WebSite', 'mainEntity has 6 items; award[] in Organization'],
        ['/about', 'FAQPage', 'mainEntity has question/answer pairs in correct language'],
        ['/insights/[slug]', 'Article', 'headline, datePublished, author.name, publisher.@id not undefined'],
        ['/team/[slug]', 'Person', 'name, jobTitle, email, sameAs (LinkedIn URL), worksFor.@id'],
    ],
    col_widths=[4.5, 5.5, 7]
)

# 1.6 Sitemap
add_heading(doc, '1.6  Sitemap', 2)
add_body(doc, 'Open http://localhost:3000/sitemap.xml and verify:')
add_bullet(doc, '/insights and /es/insights appear as <url> entries')
add_bullet(doc, 'At least one individual insight article slug appears (e.g., /insights/some-article-slug)')
add_bullet(doc, 'Each insight entry has <xhtml:link rel="alternate" hreflang="en"> and hreflang="es" child elements')
add_bullet(doc, 'No duplicate <loc> entries')

# 1.7 Build
add_heading(doc, '1.7  Build Check — Zero Regressions', 2)
add_body(doc, 'Run npm run build in the terminal. It must complete cleanly:')
add_bullet(doc, 'No TypeScript errors')
add_bullet(doc, 'No "params should be awaited" or generateMetadata warnings')
add_bullet(doc, 'No missing translation key errors')
add_bullet(doc, 'All routes pre-render without crashing')

# ── Part 2 ───────────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, 'Part 2 — Offline Discoverability Checks', 1)
add_note(doc, 'Run these after the branch is merged and deployed to production (agoralatam.com live).')

# 2.1 Google
add_heading(doc, '2.1  Google Search Operator Tests', 2)
add_body(doc, 'Run these queries directly in Google Search:')
add_table(doc,
    headers=['Query', 'What to look for'],
    rows=[
        ['site:agoralatam.com', 'Total indexed pages. /services should disappear over time; /practices should grow'],
        ['site:agoralatam.com/insights', 'Individual article URLs must appear — not just the index'],
        ['site:agoralatam.com inurl:es', 'Spanish pages (/es/*) should be indexed separately'],
        ['cache:agoralatam.com/about', 'Cached version must not show canonical: / in the source'],
        ['agoralatam.com investment arbitration Venezuela', 'Firm should appear in top results for branded + practice queries'],
    ],
    col_widths=[7.5, 9.5]
)

# 2.2 Rich Results Test
add_heading(doc, '2.2  Google Rich Results Test', 2)
add_body(doc, 'Go to https://search.google.com/test/rich-results and test these live URLs:')
add_table(doc,
    headers=['URL', 'Expected Result'],
    rows=[
        ['https://www.agoralatam.com/', 'FAQPage detected — no errors'],
        ['https://www.agoralatam.com/about', 'FAQPage detected — no errors'],
        ['Any insight article URL', 'Article detected with valid headline, author, datePublished'],
        ['Any team member profile URL', 'Person parsed (may show "ineligible for rich results" — errors must still be zero)'],
    ],
    col_widths=[7.5, 9.5]
)

# 2.3 GSC
add_heading(doc, '2.3  Google Search Console', 2)
add_body(doc, 'In Google Search Console (search.google.com/search-console):')
add_bullet(doc, 'Coverage report → confirm zero pages flagged as "Duplicate without user-selected canonical"')
add_bullet(doc, 'Sitemaps → submit https://www.agoralatam.com/sitemap.xml if not already submitted')
add_bullet(doc, 'Enhancements → FAQ → confirm FAQPage rich result detected on / and /about')
add_bullet(doc, 'International Targeting → confirm hreflang errors are zero')

# 2.4 LLMs
add_heading(doc, '2.4  LLM Discoverability Tests', 2)
add_body(doc, (
    'Use ChatGPT (with Browse), Perplexity, and Claude with web search enabled. '
    'These tests reveal whether Ágora\'s entity is known and whether pages are being cited '
    'as sources in AI answers.'
))
add_heading(doc, 'Entity Recognition Prompts', 3)
add_bullet(doc, '"Who is Ágora Latam?"')
add_bullet(doc, '"What does Ágora law firm do in Latin America?"')
add_bullet(doc, '"Is Ágora Latam ranked by Chambers?"')
add_heading(doc, 'Service / Intent Prompts', 3)
add_bullet(doc, '"Which law firms handle investment arbitration in Venezuela?"')
add_bullet(doc, '"Best law firm for cross-border M&A in Latin America"')
add_bullet(doc, '"Who are the top tax lawyers in Venezuela?"')
add_bullet(doc, '"How does foreign investment arbitration work in Latin America?"')
add_heading(doc, 'What to Look For', 3)
add_bullet(doc, 'Does the LLM cite agoralatam.com as a source?')
add_bullet(doc, 'Does it pull correct practice areas, attorney names, or firm description?')
add_bullet(doc, 'If not surfaced — these exact prompts are your highest-priority content targets from the editorial calendar')

# 2.5 Bing
add_heading(doc, '2.5  Bing Webmaster Tools', 2)
add_body(doc, 'Go to https://www.bing.com/webmasters and:')
add_bullet(doc, 'Submit https://www.agoralatam.com/sitemap.xml (separate from Google — Bing does not auto-read GSC)')
add_bullet(doc, 'Run the SEO Analyzer on / and /about — flag any canonical or hreflang errors')

# 2.6 hreflang checker
add_heading(doc, '2.6  hreflang Symmetry Check', 2)
add_body(doc, 'Go to https://hreflang.org/checker and test:')
add_bullet(doc, 'https://www.agoralatam.com/')
add_bullet(doc, 'https://www.agoralatam.com/es')
add_body(doc, 'Both pages must reference each other symmetrically. Tool will flag any one-way references or missing x-default.')

# ── Summary Table ─────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, 'Master Checklist', 1)

add_table(doc,
    headers=['#', 'Test', 'Method', 'When'],
    rows=[
        ['1', 'Canonical tags on 12 pages', 'Browser DevTools → <head>', 'Local dev'],
        ['2', 'hreflang symmetry (EN/ES/x-default)', 'Browser DevTools → <head>', 'Local dev'],
        ['3', 'FAQ accordion — visual + i18n', 'Browser visual + click test', 'Local dev'],
        ['4', '/services 301 redirect (EN + ES)', 'Browser Network tab', 'Local dev'],
        ['5', 'JSON-LD — FAQPage (homepage + about)', 'Rich Results Test (paste HTML)', 'Local dev'],
        ['6', 'JSON-LD — Article (insight pages)', 'Rich Results Test (paste HTML)', 'Local dev'],
        ['7', 'JSON-LD — Person (team profiles)', 'Rich Results Test (paste HTML)', 'Local dev'],
        ['8', 'Sitemap includes insight articles', 'Browser → /sitemap.xml', 'Local dev'],
        ['9', 'npm run build — zero errors', 'Terminal', 'Local dev'],
        ['10', 'Google site: operators', 'Google Search', 'Post-deploy'],
        ['11', 'Google Rich Results Test (live URLs)', 'search.google.com/test/rich-results', 'Post-deploy'],
        ['12', 'GSC Coverage + Sitemaps + FAQ enhancement', 'Google Search Console', 'Post-deploy'],
        ['13', 'LLM entity + service queries (3 LLMs)', 'ChatGPT Browse / Perplexity / Claude', 'Post-deploy'],
        ['14', 'Bing sitemap submission + SEO Analyzer', 'Bing Webmaster Tools', 'Post-deploy'],
        ['15', 'hreflang symmetry check (live)', 'hreflang.org/checker', 'Post-deploy'],
    ],
    col_widths=[1, 7, 5.5, 3.5]
)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    'Generated by Cursor AI · Ágora SEO & AEO Implementation · April 2026 · '
    'Branch: cleanup/seo-fixes → dev'
)
set_font(run, size=9, color=(160, 160, 160))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = 'Agora_SEO_Test_Plan.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
